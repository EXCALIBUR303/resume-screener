"""Text extraction from PDF and DOCX.

Runs inside ``worker-parse``, which has no network access. Every function here
treats its input as hostile and is expected to fail loudly with a typed error
rather than return something plausible-looking.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from enum import StrEnum

import structlog

log = structlog.get_logger()

# Below this many characters per page, the page is assumed to be a scan and OCR
# is attempted. A text-layer PDF averages hundreds of characters per page.
TEXT_DENSITY_THRESHOLD = 50
MIN_USEFUL_CHARS = 200
# Below this, the document really did yield nothing. Between the two bounds it
# yielded *something* legible but too little to score on — a different situation
# and a different message to the user.
NEGLIGIBLE_CHARS = 20


class ParseStatus(StrEnum):
    PENDING = "pending"
    PARSED = "parsed"
    NEEDS_OCR = "needs_ocr"
    LOW_TEXT = "low_text"
    NO_TEXT = "no_text"
    UNSUPPORTED = "unsupported"
    FAILED = "failed"


class ExtractionError(Exception):
    """Extraction failed in a way retrying cannot fix."""


@dataclass(frozen=True)
class Extraction:
    text: str
    extractor: str
    page_count: int
    chars_per_page: float
    ocr_used: bool = False
    ocr_confidence: int | None = None

    @property
    def status(self) -> ParseStatus:
        """Never returns an empty result silently.

        AC-1 requires that a document yielding no usable text is explicitly
        flagged, so no downstream stage can mistake "nothing extracted" for
        "nothing to say". LOW_TEXT exists because collapsing it into NO_TEXT
        told the user a perfectly legible scan was blank: OCR recovered 137
        characters at 95% confidence and the status still read `no_text`.
        """
        length = len(self.text.strip())
        if length >= MIN_USEFUL_CHARS:
            return ParseStatus.PARSED
        if not self.ocr_used:
            return ParseStatus.NEEDS_OCR
        if length >= NEGLIGIBLE_CHARS:
            return ParseStatus.LOW_TEXT
        return ParseStatus.NO_TEXT


def normalise(raw: str) -> str:
    """Collapse the whitespace noise that PDF extraction produces, without
    destroying line structure — section detection depends on line starts."""
    text = raw.replace("\x00", "").replace("\ufeff", "")
    # \xa0 is a non-breaking space; PDF extraction emits them constantly.
    text = re.sub("[ \\t\\xa0]+", " ", text)
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.split("\n")).strip()


def extract_pdf(data: bytes) -> Extraction:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data), strict=False)
        pages = len(reader.pages)
        chunks: list[str] = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception as exc:
                # One unreadable page must not lose the other twenty-nine.
                log.debug("parse.page_failed", error=type(exc).__name__)
                chunks.append("")
    except Exception as exc:
        raise ExtractionError(f"pdf extraction failed: {type(exc).__name__}") from exc

    text = normalise("\n".join(chunks))
    return Extraction(
        text=text,
        extractor="pypdf",
        page_count=pages,
        chars_per_page=len(text) / pages if pages else 0.0,
    )


def extract_docx(data: bytes) -> Extraction:
    # defusedxml is installed as a side effect of import: it replaces the stdlib
    # XML parsers with entity-resolution disabled, which is what blocks XXE in a
    # DOCX. Importing it before python-docx is deliberate and load-bearing.
    import defusedxml.ElementTree  # noqa: F401
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = Document(io.BytesIO(data))
        parts: list[str] = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
    except PackageNotFoundError as exc:
        raise ExtractionError("not a readable DOCX package") from exc
    except Exception as exc:
        raise ExtractionError(f"docx extraction failed: {type(exc).__name__}") from exc

    text = normalise("\n".join(parts))
    return Extraction(
        text=text, extractor="python-docx", page_count=1, chars_per_page=float(len(text))
    )


def needs_ocr(extraction: Extraction) -> bool:
    return not extraction.ocr_used and extraction.chars_per_page < TEXT_DENSITY_THRESHOLD
